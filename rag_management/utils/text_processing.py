# app/rag_management/utils/text_processing.py
import os
import logging
from typing import Tuple, Dict, Any, List
import re

logger = logging.getLogger(__name__)

def process_document(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text and metadata from a document
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return process_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return process_docx(file_path)
    elif file_extension in ['.txt', '.md']:
        return process_text_file(file_path)
    else:
        # Default to treating as text
        logger.warning(f"Unsupported file type {file_extension}, treating as text")
        return process_text_file(file_path)

def process_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from PDF"""
    try:
        import pypdf
        
        pdf_document = pypdf.PdfReader(file_path)
        
        # Extract metadata
        metadata = {
            "page_count": len(pdf_document.pages),
            "title": pdf_document.metadata.get('/Title', ''),
            "author": pdf_document.metadata.get('/Author', ''),
            "subject": pdf_document.metadata.get('/Subject', ''),
            "producer": pdf_document.metadata.get('/Producer', '')
        }
        
        # Extract text with page numbers
        pages_text = []
        for i, page in enumerate(pdf_document.pages):
            text = page.extract_text()
            pages_text.append({
                "page_number": i + 1,
                "text": text or ""
            })
        
        # Combine text for processing
        full_text = "\n\n".join([page["text"] for page in pages_text])
        
        # Add page information to metadata
        metadata["pages"] = pages_text
        
        return full_text, metadata
        
    except ImportError:
        logger.error("PyPDF not installed, cannot process PDF")
        return "", {"error": "PyPDF not installed"}
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {e}")
        return "", {"error": str(e)}

def process_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from DOCX"""
    try:
        import docx
        
        doc = docx.Document(file_path)
        
        # Extract metadata
        metadata = {
            "paragraph_count": len(doc.paragraphs),
        }
        
        # Extract text
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return text, metadata
        
    except ImportError:
        logger.error("python-docx not installed, cannot process DOCX")
        return "", {"error": "python-docx not installed"}
    except Exception as e:
        logger.error(f"Error processing DOCX {file_path}: {e}")
        return "", {"error": str(e)}

def process_text_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Basic metadata
        metadata = {
            "line_count": text.count('\n') + 1,
            "character_count": len(text),
            "word_count": len(text.split())
        }
        
        return text, metadata
        
    except Exception as e:
        logger.error(f"Error processing text file {file_path}: {e}")
        return "", {"error": str(e)}

def split_text_into_chunks(text: str, metadata: Dict[str, Any] = None, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks for processing
    
    Args:
        text: Text to split
        metadata: Optional metadata about the document
        chunk_size: Target size of each chunk in characters
        chunk_overlap: Overlap between chunks in characters
        
    Returns:
        List of dictionaries containing content and metadata for each chunk
    """
    if not text:
        return []
    
    # Simple splitting by paragraphs then recombining to meet chunk size
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    current_chunk_metadata = {}
    
    # Track the page we're currently processing
    current_page = 1
    page_map = {}
    
    # If we have page information in the metadata, create a character-to-page mapping
    if metadata and "pages" in metadata:
        char_count = 0
        for page_info in metadata["pages"]:
            page_num = page_info["page_number"]
            page_text = page_info["text"]
            page_length = len(page_text)
            
            # Map character ranges to page numbers
            page_map[(char_count, char_count + page_length)] = page_num
            char_count += page_length + 2  # +2 for the newlines we added between pages
    
    # Process paragraphs
    char_position = 0
    for para in paragraphs:
        # Find which page this paragraph belongs to based on character position
        if page_map:
            for (start, end), page in page_map.items():
                if start <= char_position < end:
                    current_page = page
                    break
        
        # If adding this paragraph would exceed the chunk size, save the current chunk and start a new one
        if current_chunk and len(current_chunk) + len(para) + 2 > chunk_size:
            # Save the current chunk
            chunks.append({
                "content": current_chunk,
                "metadata": {
                    "page_number": current_page,
                    **current_chunk_metadata
                }
            })
            
            # Start a new chunk with overlap
            # Find the last complete sentence or paragraph boundary within the overlap window
            overlap_start = max(0, len(current_chunk) - chunk_overlap)
            if overlap_start > 0:
                # Try to find sentence boundaries
                sentences = re.split(r'(?<=[.!?])\s+', current_chunk[overlap_start:])
                if len(sentences) > 1:
                    # We found sentence boundaries, start from the beginning of the second-to-last sentence
                    last_sentence_pos = current_chunk[overlap_start:].find(sentences[-2])
                    if last_sentence_pos != -1:
                        current_chunk = current_chunk[overlap_start + last_sentence_pos:]
                    else:
                        current_chunk = current_chunk[overlap_start:]
                else:
                    current_chunk = current_chunk[overlap_start:]
            else:
                current_chunk = ""
        
        # Add the paragraph to the current chunk
        if current_chunk:
            current_chunk += "\n\n" + para
        else:
            current_chunk = para
        
        # Update character position
        char_position += len(para) + 2  # +2 for newline chars
    
    # Add the final chunk if it's not empty
    if current_chunk:
        chunks.append({
            "content": current_chunk,
            "metadata": {
                "page_number": current_page,
                **current_chunk_metadata
            }
        })
    
    return chunks